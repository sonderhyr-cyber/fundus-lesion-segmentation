# -*- coding: utf-8 -*-
"""Generate the plain-text Word report for the HRDecoder SE experiments."""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = r"D:\dataset\lesion_segmentation\00_总报告_HRDecoder_SE专项实验报告.docx"

doc = Document()

# Default fonts (Chinese-friendly)
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    return p


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(10)
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    p.paragraph_format.space_before = Pt(6)
    return p


def para(text):
    return doc.add_paragraph(text)


def add_figure(path, caption):
    p = doc.add_paragraph()
    p.alignment = 1  # centered
    run = p.add_run()
    run.add_picture(path, width=Inches(5.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = 1
    for r in cap.runs:
        r.font.size = Pt(9)
    return p


title("HRDecoder SE病灶专项实验报告")
para("说明：本报告为纯文字汇总，无排版装饰。")

h1("一、实验背景与设定")
h2("1.1 项目背景")
para("项目基于HRDecoder（MICCAI 2024，Ding et al.，arXiv:2411.03976），"
     "使用MMSegmentation 0.16.0 / mmcv 1.7.1 / Python 3.8.5 / PyTorch 2.0.1。")
para("实验数据为DDR分割子集共757张图（原始划分train 383 / val 149 / test 225），"
     "四类病灶标签为EX（硬性渗出）、HE（出血）、SE（软性渗出）、MA（微动脉瘤）。"
     "标签以labels/{split}/{EX,HE,SE,MA}下的tif原始标注及融合单通道png"
     "（EX=1，HE=2，SE=3，MA=4）形式使用。")
h2("1.2 实验设计说明")
para("数据集按6:1:3重新划分（train 454 / val 76 / test 227）为自定义实验设计，"
     "并非论文原始划分。")
para("板块二（SE copy-paste增强、SE滑动窗口过采样）与板块三（L_Seg中SE类别加权）"
     "均为在论文框架基础上自行设计的扩展实验，不是论文原文方法。"
     "论文官方训练流程为随机缩放、随机裁剪、随机翻转、随机旋转等标准增强，"
     "即本实验baseline使用的流程。")
para("所有实验为单次训练，随机种子固定为42，训练40000 iter，"
     "用val集（76张）选择checkpoint，test集（227张）仅在最终统一评估，"
     "全程未用于调参。")
h2("1.3 环境与硬件")
para("4台RTX 4090服务器并行训练；服务器端Linux，conda Python 3.8.10，"
     "PyTorch 2.0.1+cu117，MMCV 1.7.1，MMSegmentation 0.16.0（项目自带源码）。")

h1("二、板块一：DDR数据集6:1:3重新划分")
h2("2.1 划分方法")
para("将原始train/val/test共757张图合并为一个池子，按SE标签分层抽样："
     "档A为不含SE病灶的图；档B为含SE且SE像素比例低于SE阳性图中位数的图；"
     "档C为含SE且SE像素比例不低于中位数的图。")
para("各档内部按6:1:3比例分配（最大余数法），再将三档的train/val/test合并，"
     "随机种子固定为42。")
para("输出：images/{train,val,test}与labels/{train,val,test}目录（符号链接）、"
     "train.txt/val.txt/test.txt文件列表、split_assignment.json逐图记录、"
     "split_stats.csv与split_report.md统计报告。")
h2("2.2 划分结果")
para("最终规模：train 454张，val 76张，test 227张，合计757张。")
para("SE图片占比：train 31.50%，val 31.58%，test 31.72%，整体31.57%，"
     "三组与整体偏差不超过0.15个百分点（验收容差为正负5个百分点，通过）。")
para("SE像素占比：train 0.0566%，val 0.0502%，test 0.0482%，整体0.0534%。")
para("对照类别图片占比（EX/HE/MA）：train 66.74%/80.62%/74.45%，"
     "val 60.53%/82.89%/77.63%，test 60.35%/75.77%/76.21%，"
     "整体64.20%/79.39%/75.30%，波动均在正负10个百分点内，如实记录。")
h2("2.3 数据质量说明")
para("757张图图片与标签完整配对，无缺失文件。")
para("141张图的融合png与原始tif存在微小像素差异，共160处"
     "（EX 76处、HE 78处、SE 6处；中位相对差小于0.6%，SE最大差异52像素，MA无差异）。"
     "训练使用融合png，影响可忽略；发现的问题仅报告，未做删改。")
h2("2.4 baseline对照")
para("同一配置在原官方划分上复跑的测试mIoU为30.81%；"
     "新划分下baseline测试mIoU为42.20%。"
     "论文官方划分DDR mIoU约32%，新测试集病灶构成与官方测试集不同"
     "（新测试集EX/HE像素占比更低）且训练数据多18.5%，故数字偏高属正常，"
     "已核查非划分引入的难例集中偏差。")

h1("三、板块二：SE数据增强实验")
h2("3.1 方法")
para("② +SE copy-paste：从训练集143张SE阳性图提取588个SE连通域病灶块"
     "（图像+标签同步裁剪），建立病灶库；训练时以概率0.5将最多2个病灶块"
     "粘贴到其他图像的视网膜FOV内（灰度阈值加腐蚀生成FOV，避免贴到背景），"
     "同时更新GT标签。")
para("③ +SE滑动窗口过采样：将HR分支的原随机裁剪改为滑动窗口加权采样，"
     "窗口1024x1024、步长512，按窗口内SE像素占比加权"
     "（权重=1+20乘归一化SE占比），每次采样4个窗口。")
para("④ ②+③组合。")
h2("3.2 测试集结果（IoU%/F1%/AUPR%）")
para("baseline：mIoU 42.20，mF1 58.45，mAUPR 60.59；"
     "EX 51.56/68.04/72.66，HE 43.13/60.26/63.82，SE 49.98/66.65/69.60，"
     "MA 24.11/38.85/36.27。")
para("② copy-paste：mIoU 39.67，mF1 56.08，mAUPR 58.77；"
     "EX 50.54/67.14/71.39，HE 39.62/56.75/61.35，SE 44.22/61.32/67.67，"
     "MA 24.30/39.09/34.68。")
para("③ 滑窗过采样：mIoU 40.07，mF1 56.38，mAUPR 58.39；"
     "EX 51.77/68.22/72.75，HE 38.22/55.31/59.95，SE 46.30/63.29/66.42，"
     "MA 24.01/38.72/34.44。")
para("④ 组合：mIoU 40.79，mF1 57.16，mAUPR 60.13；"
     "EX 51.46/67.95/72.55，HE 42.29/59.44/64.33，SE 45.28/62.34/68.43，"
     "MA 24.15/38.91/35.23。")
para("val最佳mIoU：baseline 39.4%（32k iter），② 39.7%（32k），"
     "③ 39.6%（20k），④ 40.4%（40k）。")
h2("3.3 结论")
para("三组增强在val上均不低于baseline，但测试集SE IoU均未超过baseline的49.98，"
     "整体mIoU下降1.4至2.5个百分点。")
para("②④（copy-paste系）显著提升SE精确率PPV（83.87%与81.49%，"
     "baseline为72.78%），但召回率下降（48.33%与50.47%，baseline为61.47%），"
     "SE IoU净下降，模型预测SE更保守。")
para("③滑窗在val有明显增益（39.6%），但测试集未转化（SE 46.30），"
     "存在val/test泛化差距。")
para("本轮为单种子结果，作为可复现的对照数据保留，不做美化。")

h1("四、板块三：损失函数加权实验")
h2("4.1 方法")
para("在BinaryLoss（Dice）中引入类别权重class_weight=[EX,HE,SE,MA]，"
     "作用于L_Seg的逐类别求和。")
para("实验组：1x（baseline）、1.5x、2x、3x，"
     "以及按训练集像素频率倒数计算的自适应权重"
     "（归一化均值=1）：[0.236,0.154,1.027,2.584]。")
h2("4.2 测试集结果")
para("1x：mIoU 42.20，SE IoU 49.98，HE 43.13，EX 51.56，MA 24.11。")
para("1.5x：mIoU 40.90，SE 45.68，HE 43.96，EX 49.47，MA 24.48。")
para("2x：mIoU 40.31，SE 44.77，HE 42.02，EX 51.00，MA 23.45。")
para("3x：mIoU 38.50，SE 42.61，HE 37.81，EX 49.79，MA 23.82。")
para("自适应：mIoU 38.76，SE 48.52，HE 32.58，EX 50.62，MA 23.31。")
add_figure(r"D:\dataset\lesion_segmentation\server_checkpoints\03_板块三_SE权重mIoU曲线.png",
           "图1：SE权重与测试mIoU曲线（1x为baseline，灰色虚线为自适应组mIoU）")
h2("4.3 结论")
para("SE权重从1x增至3x，整体mIoU单调下降（42.20到40.90到40.31到38.50），"
     "SE自身IoU也未提升；1.5x为相对最优组，其HE IoU 43.96为各组最高。")
para("自适应逆频率权重主要牺牲HE（32.58，较baseline下降10.55个百分点），"
     "SE为48.52接近baseline，整体不划算。")
para("推测原因：HRDecoder的L_Seg为逐类别二值Dice之和，SE标注稀疏，"
     "提高SE权重在放大SE监督的同时压制EX/HE学习，未能带来SE收益。")

h1("五、交付物")
para("服务器端（实例已关机，数据保留在实例磁盘）：")
para("/root/HRDecoder/data/DDR_split_613/：数据目录，含train.txt、val.txt、"
     "test.txt、split_stats.csv、split_report.md、split_assignment.json、"
     "se_lesion_library.pkl、baseline_report.md。")
para("/root/HRDecoder/ddr_split.py：数据集划分脚本。")
para("/root/HRDecoder/build_se_library.py：SE病灶库构建脚本。")
para("/root/HRDecoder/mmseg/datasets/pipelines/se_aug.py：SE增强实现。")
para("/root/HRDecoder/mmseg/models/segmentors/HRDecoder.py："
     "修改后的模型（se_slide滑窗采样模式）。")
para("/root/HRDecoder/configs/lesion/：8个实验配置（baseline、"
     "se_copypaste、se_slide、se_copypaste_slide、sew15、sew20、sew30、"
     "sew_adaptive）。")
para("/root/HRDecoder/work_dirs/：各实验子目录，含训练日志与checkpoint。")
para("本地工程文件（D:\\dataset\\lesion_segmentation\\，已按板块改名）：")
para("00_总报告_HRDecoder_SE专项实验报告.docx：本报告。")
para("00_工具_SSH连接脚本_ssh_remote.py、00_工具_多机管理脚本_remote.ps1、"
     "00_工具_服务器清单_servers.json、00_工具_四机巡检脚本_status_all.ps1。")
para("01_板块一_数据集划分脚本_ddr_split.py、"
     "01_板块一_baseline实验配置.py、"
     "01_板块一_划分统计报告_baseline_report.md。")
para("02_板块二_病灶库构建脚本_build_se_library.py、"
     "02_板块二_增强实现_se_aug.py、"
     "02_板块二_模型滑窗采样实现_HRDecoder.py、"
     "02_板块二_pipeline注册文件_pipelines_init.py、"
     "02_板块二_实验2_copyPaste配置.py、"
     "02_板块二_实验3_滑窗配置.py、"
     "02_板块二_实验4_组合配置.py。")
para("03_板块三_实验1.5x配置.py、03_板块三_实验2x配置.py、"
     "03_板块三_实验3x配置.py、03_板块三_实验自适应配置.py、"
     "03_板块三_绘图脚本_plot_results.py。")
para("04_板块二三汇总报告_module2_3_report.md、"
     "04_工具_报告生成脚本_build_report_docx.py。")
para("本地JSON工程日志（D:\\dataset\\lesion_segmentation\\实验日志_json\\）：")
para("01_板块一_baseline.json；02_板块二_实验2_copyPaste.json及前半段"
     "0-2300iter.json（续跑日志从2000开始，前半段日志补齐0-2300）；"
     "02_板块二_实验3_滑窗.json；02_板块二_实验4_组合.json；"
     "03_板块三_1.5x.json；03_板块三_2x.json；03_板块三_3x.json；"
     "03_板块三_自适应.json。每份日志为JSONL格式，"
     "含400个训练记录点（每100 iter一条）与10次val评估（每4000 iter一次）。")
para("本地checkpoint备份（D:\\dataset\\lesion_segmentation\\server_checkpoints\\）：")
para("01_板块一_baseline_best_iter32000.pth、"
     "02_板块二_实验2_copyPaste_best_iter32000.pth、"
     "02_板块二_实验3_滑窗_best_iter20000.pth、"
     "02_板块二_实验4_组合_best_iter40000.pth、"
     "03_板块三_1.5x_best_iter20000.pth、"
     "03_板块三_2x_best_iter20000.pth、"
     "03_板块三_3x_best_iter20000.pth、"
     "03_板块三_自适应_best_iter40000.pth、"
     "03_板块三_SE权重mIoU曲线.png。")

h1("六、结论与后续建议")
para("本轮实验显示：在固定6:1:3划分、单种子42、40000 iter条件下，"
     "基于论文框架自行设计的SE copy-paste、滑窗过采样与SE类别加权"
     "均未在测试集上带来SE IoU提升；copy-paste提升SE精确率但损失召回，"
     "loss加权随权重增大整体下降。")
para("后续可选方向：1）对关键对比组补跑种子123、2024做三次重复平均；"
     "2）调整copy-paste参数（病灶大小过滤、粘贴概率、粘贴数量）重新迭代；"
     "3）在val上分析SE漏检与误检模式后再设计增强策略。")

doc.save(OUT)
print("saved:", OUT)
